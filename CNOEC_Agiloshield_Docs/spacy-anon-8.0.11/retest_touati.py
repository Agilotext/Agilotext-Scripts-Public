#!/usr/bin/env python3
"""
retest_touati.py — Script de re-test des 7 fichiers du client Touati

Ce script :
1. Extrait le texte brut des fichiers DOCX originaux (NON anonymisés)
2. Applique la version active du processeur d'anonymisation
3. Analyse les résultats : compte les tags par type, identifie les résidus suspects
4. Compare avec les résultats de la baseline v8.0.7 (si disponible)
5. Génère un rapport Markdown

Usage :
    python retest_touati.py --input-dir /chemin/vers/originaux --output-dir ./results

Prérequis :
    pip install python-docx spacy
    python -m spacy download fr_core_news_lg
"""

import argparse
import os
import re
import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import spacy
    HAS_SPACY = True
except ImportError:
    HAS_SPACY = False

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from anon.anon_processor import AnonProcessor
try:
    from anon.version import version as ANON_VERSION
except Exception:
    ANON_VERSION = "unknown"


TAG_PATTERN = re.compile(r"\[(PERSON|ORGANIZATION|LOCATION|BIRTH_DATE|BIRTH_PLACE|"
                         r"EMAIL|URL|PHONE|IBAN|BIC|RIB|VAT|SIRET|SIREN|NIR|"
                         r"URSSAF_ID|FISCAL_ID|APE|POSTAL_CODE|DATE)\]")


def extract_text_from_docx(filepath: str) -> str:
    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def count_tags(text: str) -> dict:
    counts = {}
    for m in TAG_PATTERN.finditer(text):
        tag = m.group(1)
        counts[tag] = counts.get(tag, 0) + 1
    return counts


def find_suspect_residuals(anonymized_text: str) -> list:
    """Cherche des résidus potentiellement non anonymisés."""
    suspects = []

    email_re = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
    for m in email_re.finditer(anonymized_text):
        suspects.append(("EMAIL_RESIDU", m.group(0), m.start()))

    phone_re = re.compile(r"(?<!\d)0[67][\s.\-]?\d{2}[\s.\-]?\d{2}[\s.\-]?\d{2}[\s.\-]?\d{2}(?!\d)")
    for m in phone_re.finditer(anonymized_text):
        suspects.append(("PHONE_RESIDU", m.group(0), m.start()))

    name_re = re.compile(r"(?<!\[)(?<![A-Za-z])([A-ZÀ-Ÿ][a-zà-ÿ]+\s+[A-ZÀ-Ÿ]{2,})(?!\])")
    for m in name_re.finditer(anonymized_text):
        word = m.group(1)
        lower_words = word.lower().split()
        stopwords = {"de", "la", "le", "du", "les", "des", "en", "et", "au", "aux",
                     "rue", "avenue", "boulevard", "total", "classe"}
        if not any(w in stopwords for w in lower_words):
            suspects.append(("NAME_RESIDU", m.group(1), m.start()))

    return suspects


def process_file(filepath: str, nlp) -> dict:
    filename = os.path.basename(filepath)
    print(f"  Traitement de {filename}...")

    text = extract_text_from_docx(filepath)
    anonymized = AnonProcessor.do_anon_text(text, nlp, cell_mode=False)
    tags = count_tags(anonymized)
    residuals = find_suspect_residuals(anonymized)

    return {
        "filename": filename,
        "original_length": len(text),
        "anonymized_length": len(anonymized),
        "tags": tags,
        "total_tags": sum(tags.values()),
        "residuals": residuals,
        "residual_count": len(residuals),
        "anonymized_text": anonymized,
    }


def generate_report(results: list, output_dir: str):
    report_path = os.path.join(output_dir, "RETEST_TOUATI_RESULTS.md")
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    lines = [
        f"# Re-test Anonymisation — Touati v{ANON_VERSION}",
        f"",
        f"**Date :** {now}  ",
        f"**Version processeur :** v{ANON_VERSION}  ",
        f"**Fichiers testés :** {len(results)}",
        f"",
        f"---",
        f"",
        f"## Résumé",
        f"",
        f"| Fichier | Taille (chars) | Tags appliqués | Résidus suspects |",
        f"|---------|---------------|----------------|-----------------|",
    ]

    for r in results:
        lines.append(
            f"| {r['filename'][:30]} | {r['original_length']} | {r['total_tags']} | {r['residual_count']} |"
        )

    lines.extend([
        f"",
        f"---",
        f"",
        f"## Détail par fichier",
    ])

    for r in results:
        lines.extend([
            f"",
            f"### {r['filename']}",
            f"",
            f"**Tags appliqués :**",
            f"",
        ])

        if r["tags"]:
            lines.append("| Tag | Count |")
            lines.append("|-----|-------|")
            for tag, count in sorted(r["tags"].items(), key=lambda x: -x[1]):
                lines.append(f"| [{tag}] | {count} |")
        else:
            lines.append("Aucun tag appliqué.")

        if r["residuals"]:
            lines.extend([
                f"",
                f"**Résidus suspects ({r['residual_count']}) :**",
                f"",
                f"| Type | Valeur | Position |",
                f"|------|--------|----------|",
            ])
            for rtype, val, pos in r["residuals"][:20]:
                lines.append(f"| {rtype} | `{val[:40]}` | {pos} |")
        else:
            lines.extend([f"", f"Aucun résidu suspect détecté."])

    lines.extend([
        f"",
        f"---",
        f"",
        f"## Comparaison avec baseline v8.0.7",
        f"",
        f"| Critère | v8.0.7 (baseline) | v{ANON_VERSION} |",
        f"|---------|----------------------|---------------|",
        f"| Faux positifs | ~25 | À remplir |",
        f"| Fuites (noms) | ~15 | À remplir |",
        f"| Fuites (villes) | ~8 | À remplir |",
        f"| Note globale | 5/10 | À remplir |",
    ])

    with open(report_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"\nRapport généré : {report_path}")

    json_path = os.path.join(output_dir, "retest_results.json")
    export = []
    for r in results:
        export.append({
            "filename": r["filename"],
            "original_length": r["original_length"],
            "anonymized_length": r["anonymized_length"],
            "tags": r["tags"],
            "total_tags": r["total_tags"],
            "residual_count": r["residual_count"],
            "residuals": [(t, v, p) for t, v, p in r["residuals"]],
        })
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(export, f, ensure_ascii=False, indent=2)
    print(f"Données JSON : {json_path}")


def main():
    parser = argparse.ArgumentParser(description=f"Re-test anonymisation Touati v{ANON_VERSION}")
    parser.add_argument("--input-dir", required=True, help="Dossier contenant les DOCX originaux (non anonymisés)")
    parser.add_argument("--output-dir", default="./results", help="Dossier de sortie pour le rapport")
    parser.add_argument("--model", default="fr_core_news_lg", help="Modèle SpaCy à utiliser")
    args = parser.parse_args()

    if not HAS_DOCX:
        print("ERREUR : python-docx n'est pas installé. Installez-le avec : pip install python-docx")
        sys.exit(1)
    if not HAS_SPACY:
        print("ERREUR : spacy n'est pas installé. Installez-le avec : pip install spacy")
        sys.exit(1)

    os.makedirs(args.output_dir, exist_ok=True)

    print(f"Chargement du modèle SpaCy ({args.model})...")
    nlp = spacy.load(args.model)
    nlp.max_length = 2_000_000

    docx_files = sorted([
        os.path.join(args.input_dir, f)
        for f in os.listdir(args.input_dir)
        if f.endswith(".docx") and not f.startswith("~")
    ])

    if not docx_files:
        print(f"Aucun fichier DOCX trouvé dans {args.input_dir}")
        sys.exit(1)

    print(f"\n{len(docx_files)} fichiers DOCX trouvés.\n")

    results = []
    for fp in docx_files:
        result = process_file(fp, nlp)
        results.append(result)

        anon_path = os.path.join(args.output_dir, os.path.basename(fp).replace(".docx", ".ANON.txt"))
        with open(anon_path, "w", encoding="utf-8") as f:
            f.write(result["anonymized_text"])

    generate_report(results, args.output_dir)


if __name__ == "__main__":
    main()
